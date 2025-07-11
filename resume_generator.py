import os
from jinja2 import Environment, FileSystemLoader
import pdfkit
import streamlit as st
from typing import Dict
import tempfile
from docx import Document
import docx.shared

class ResumeGenerator:
    """Handle resume generation using templates."""
    
    def __init__(self):
        """Initialize the template environment."""
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        self.env = Environment(loader=FileSystemLoader(template_dir))
        
        # Configure pdfkit to use wkhtmltopdf
        self.pdf_options = {
            'page-size': 'Letter',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in',
            'encoding': "UTF-8",
            'no-outline': None,
            'quiet': ''
        }
    
    def render_html(self, data: Dict) -> str:
        """Render the resume template with the provided data."""
        template = self.env.get_template('resume_template.html')
        return template.render(**data)
    
    def generate_pdf(self, html_content: str) -> bytes:
        """Convert HTML content to PDF."""
        try:
            # Create a temporary HTML file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.html', mode='w', encoding='utf-8') as f:
                f.write(html_content)
                temp_html = f.name
            
            # Create a temporary PDF file
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            
            # Generate PDF
            pdfkit.from_file(temp_html, temp_pdf, options=self.pdf_options)
            
            # Read the PDF content
            with open(temp_pdf, 'rb') as f:
                pdf_content = f.read()
            
            # Clean up temporary files
            os.unlink(temp_html)
            os.unlink(temp_pdf)
            
            return pdf_content
            
        except Exception as e:
            st.error(f"Error generating PDF: {str(e)}")
            st.info("Please ensure wkhtmltopdf is installed on your system.")
            return None
    
    def format_resume_data(self, form_data: Dict) -> Dict:
        """Format the form data for template rendering."""
        return {
            'personal_info': form_data['personal_info'],
            'education': form_data['education'],
            'experience': form_data['experience'],
            'skills': form_data['skills'],
            'certifications': form_data['certifications']
        }
    
    def generate_docx(self, resume_data: Dict) -> Document:
        """Generate a formatted DOCX resume using python-docx."""
        doc = Document()
        
        # Add personal info section
        name = resume_data['personal_info'].get('full_name', '')
        doc.add_heading(name, 0)
        
        # Contact info
        contact_info = []
        if resume_data['personal_info'].get('email'):
            contact_info.append(resume_data['personal_info']['email'])
        if resume_data['personal_info'].get('phone'):
            contact_info.append(resume_data['personal_info']['phone'])
        if resume_data['personal_info'].get('location'):
            contact_info.append(resume_data['personal_info']['location'])
        if resume_data['personal_info'].get('linkedin'):
            contact_info.append(resume_data['personal_info']['linkedin'])
        
        doc.add_paragraph(' | '.join(contact_info))
        doc.add_paragraph()  # Add spacing
        
        # Summary section
        if resume_data['personal_info'].get('summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(resume_data['personal_info']['summary'])
            doc.add_paragraph()
        
        # Experience section
        if resume_data.get('experience'):
            doc.add_heading('Experience', level=1)
            for exp in resume_data['experience']:
                p = doc.add_paragraph()
                p.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}").bold = True
                p.add_run(f"\n{exp.get('start_date', '')} - {exp.get('end_date', '')}")
                
                # Add responsibilities as bullet points
                for resp in exp.get('responsibilities', []):
                    doc.add_paragraph(resp, style='List Bullet')
            doc.add_paragraph()
        
        # Education section
        if resume_data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in resume_data['education']:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')}").bold = True
                p.add_run(f"\n{edu.get('institution', '')}, {edu.get('start_year', '')} - {edu.get('end_year', '')}")
            doc.add_paragraph()
        
        # Skills section
        doc.add_heading('Skills', level=1)
        if resume_data.get('skills', {}).get('technical'):
            p = doc.add_paragraph()
            p.add_run('Technical Skills: ').bold = True
            p.add_run(', '.join(resume_data['skills']['technical']))
        
        if resume_data.get('skills', {}).get('soft'):
            p = doc.add_paragraph()
            p.add_run('Soft Skills: ').bold = True
            p.add_run(', '.join(resume_data['skills']['soft']))
        doc.add_paragraph()
        
        # Certifications section
        if resume_data.get('certifications'):
            doc.add_heading('Certifications', level=1)
            for cert in resume_data['certifications']:
                p = doc.add_paragraph()
                p.add_run(f"{cert.get('name', '')}").bold = True
                p.add_run(f"\n{cert.get('issuer', '')}, {cert.get('year', '')}")
        
        # Apply some basic styling
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = docx.shared.Pt(11)
        
        # Add page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = docx.shared.Inches(0.5)
            section.bottom_margin = docx.shared.Inches(0.5)
            section.left_margin = docx.shared.Inches(0.5)
            section.right_margin = docx.shared.Inches(0.5)
        
        return doc
